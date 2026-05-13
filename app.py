import streamlit as st
import docx
import pandas as pd
import io
import os
import tempfile
import subprocess
import shutil
from zipfile import BadZipFile

st.set_page_config(page_title="Word to Excel Converter", page_icon="📝", layout="wide")

# Custom CSS for a more premium look
st.markdown("""
    <style>
    .main {
        background-color: #f8f9fa;
    }
    .stButton>button {
        width: 100%;
        border-radius: 5px;
        height: 3em;
        background-color: #007bff;
        color: white;
    }
    .stDownloadButton>button {
        width: 100%;
        border-radius: 5px;
        height: 3em;
        background-color: #28a745;
        color: white;
    }
    </style>
    """, unsafe_allow_html=True)

st.title("📝 Spec Table Extractor")
st.markdown("""
    **Only extracts tables and ignores regular text.**  
    Upload any Word document with mixed content, and I will find all tables for you.
""")

uploaded_file = st.file_uploader("Choose a Word file", type=["docx", "doc"])

def convert_doc_to_docx_with_word(doc_bytes):
    try:
        import pythoncom
        import win32com.client
    except ImportError as exc:
        raise RuntimeError(
            "Word-based conversion requires pywin32."
        ) from exc

    with tempfile.TemporaryDirectory() as tmp_dir:
        input_doc = os.path.join(tmp_dir, "input.doc")
        output_docx = os.path.join(tmp_dir, "output.docx")

        with open(input_doc, "wb") as f:
            f.write(doc_bytes)

        word = None
        document = None
        pythoncom.CoInitialize()
        try:
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0

            document = word.Documents.Open(input_doc, ReadOnly=True)
            # 16 = wdFormatDocumentDefault (.docx)
            document.SaveAs2(output_docx, FileFormat=16)
        except Exception as exc:
            raise RuntimeError(
                "Failed to convert .doc automatically. Please ensure Microsoft Word is installed and can open this file."
            ) from exc
        finally:
            if document is not None:
                document.Close(False)
            if word is not None:
                word.Quit()
            pythoncom.CoUninitialize()

        if not os.path.exists(output_docx):
            raise RuntimeError("Automatic .doc conversion did not produce a .docx file.")

        with open(output_docx, "rb") as f:
            return f.read()

def convert_doc_to_docx_with_libreoffice(doc_bytes):
    soffice_path = shutil.which("soffice")
    if not soffice_path:
        raise RuntimeError("LibreOffice (soffice) is not available in the current environment.")

    with tempfile.TemporaryDirectory() as tmp_dir:
        input_doc = os.path.join(tmp_dir, "input.doc")
        output_docx = os.path.join(tmp_dir, "input.docx")

        with open(input_doc, "wb") as f:
            f.write(doc_bytes)

        result = subprocess.run(
            [
                soffice_path,
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                tmp_dir,
                input_doc,
            ],
            capture_output=True,
            text=True,
            check=False,
        )

        if result.returncode != 0 or not os.path.exists(output_docx):
            raise RuntimeError(
                "LibreOffice conversion failed."
            )

        with open(output_docx, "rb") as f:
            return f.read()

def convert_doc_to_docx_bytes(doc_bytes):
    conversion_errors = []

    if os.name == "nt":
        try:
            return convert_doc_to_docx_with_word(doc_bytes)
        except RuntimeError as e:
            conversion_errors.append(f"Word conversion failed: {e}")

    try:
        return convert_doc_to_docx_with_libreoffice(doc_bytes)
    except RuntimeError as e:
        conversion_errors.append(f"LibreOffice conversion failed: {e}")

    raise RuntimeError(
        "Automatic .doc conversion is unavailable in this environment. "
        "Needs either Microsoft Word (Windows) or LibreOffice (soffice). "
        + " | ".join(conversion_errors)
    )

def deduplicate_columns(columns):
    new_columns = []
    seen = {}
    for col in columns:
        if col in seen:
            seen[col] += 1
            new_columns.append(f"{col}_{seen[col]}")
        else:
            seen[col] = 0
            new_columns.append(col)
    return new_columns

if uploaded_file is not None:
    try:
        file_ext = os.path.splitext(uploaded_file.name)[1].lower()
        base_name = os.path.splitext(uploaded_file.name)[0]

        if file_ext == ".doc":
            st.info("Detected .doc file. Converting to .docx automatically...")
            converted_docx_bytes = convert_doc_to_docx_bytes(uploaded_file.getvalue())
            doc = docx.Document(io.BytesIO(converted_docx_bytes))
            st.success("Conversion complete. Processing tables now.")
        else:
            doc = docx.Document(uploaded_file)

        if not doc.tables:
            st.error("No tables found in the uploaded document.")
        else:
            st.success(f"Found {len(doc.tables)} table(s)!")
            
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                for i, table in enumerate(doc.tables):
                    data = []
                    for row in table.rows:
                        # Handle potential merged cells by getting unique text in each cell
                        row_data = [cell.text.strip() for cell in row.cells]
                        data.append(row_data)
                    
                    df = pd.DataFrame(data)
                    
                    if not df.empty:
                        # Extract headers and deduplicate them
                        headers = [str(c).strip() for c in df.iloc[0]]
                        df.columns = deduplicate_columns(headers)
                        df = df[1:]
                    
                    sheet_name = f"Table_{i+1}"
                    df.to_excel(writer, sheet_name=sheet_name, index=False)
                    
                    with st.expander(f"Preview Table {i+1}"):
                        st.dataframe(df)
            
            # Finalize the data
            xlsx_data = output.getvalue()
            
            # Determine download filename
            download_filename = f"{base_name}.xlsx"
            
            st.info(f"Ready to download: **{download_filename}**")

            st.download_button(
                label="📥 Download Excel File",
                data=xlsx_data,
                file_name=download_filename,
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
            
    except BadZipFile:
        st.error("The uploaded file is not a valid .docx file. Please export it as .docx and try again.")
    except RuntimeError as e:
        st.error(str(e))
    except Exception as e:
        st.error(f"An error occurred: {e}")

st.divider()
st.info("Tip: You can drag and drop your file directly into the box above.")
